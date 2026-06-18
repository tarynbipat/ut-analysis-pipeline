"""Convert the markdown design report to a .docx file."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def md_to_docx(md_path: str, docx_path: str) -> None:
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    table_lines: list[str] = []
    in_blockquote = False
    blockquote_text = ""

    def flush_blockquote():
        nonlocal in_blockquote, blockquote_text
        if blockquote_text.strip():
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(blockquote_text.strip())
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
        in_blockquote = False
        blockquote_text = ""

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        # Parse markdown table
        rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.strip("|").split("|")]
            rows.append(cells)
        # Remove separator row (---) 
        rows = [r for r in rows if not all(set(c) <= set("- :") for c in r)]
        if not rows:
            table_lines = []
            return
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell
                    if i == 0:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        doc.add_paragraph()
        table_lines = []

    for line in lines:
        line = line.rstrip("\n")

        # Table row
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if in_blockquote:
                flush_blockquote()
            table_lines.append(line)
            continue
        elif table_lines:
            flush_table()

        # Blockquote
        if line.strip().startswith(">"):
            in_blockquote = True
            text = re.sub(r"^>\s*", "", line.strip())
            blockquote_text += text + " "
            continue
        elif in_blockquote:
            flush_blockquote()

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        elif line.strip().startswith("- ") or line.strip().startswith("| "):
            # Bullet point
            text = re.sub(r"^[\s\-\|]+", "", line)
            doc.add_paragraph(text, style="List Bullet")
        elif line.strip():
            # Regular paragraph - handle inline bold/italic
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    # Flush remaining
    if in_blockquote:
        flush_blockquote()
    if table_lines:
        flush_table()

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    md_to_docx(
        "project_data/data/reports/design_report.md",
        "project_data/data/reports/Embr_UT_Design_Report.docx",
    )
